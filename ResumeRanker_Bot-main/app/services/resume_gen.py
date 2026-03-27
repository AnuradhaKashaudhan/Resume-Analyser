"""
resume_gen.py – Generate downloadable PDF and DOCX resume files.

Libraries:
  • reportlab  → PDF generation
  • python-docx → DOCX generation
"""

from __future__ import annotations

from pathlib import Path

from app.utils.helpers import setup_logging, TEMP_DIR

logger = setup_logging("drcode.resume_gen")


def generate_pdf(text: str, filename: str = "improved_resume.pdf") -> Path:
    """
    Create a professionally-formatted PDF from plain text.

    Args:
        text:     The improved resume text.
        filename: Output filename.

    Returns:
        Path to the generated PDF file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    )

    output_path = TEMP_DIR / filename
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=HexColor("#1a1a2e"),
        spaceAfter=6,
        spaceBefore=12,
        fontName="Helvetica-Bold",
    )

    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        textColor=HexColor("#333333"),
        fontName="Helvetica",
    )

    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontSize=18,
        textColor=HexColor("#0f3460"),
        spaceAfter=8,
        fontName="Helvetica-Bold",
    )

    flowables = []

    # Title
    flowables.append(Paragraph("Improved Resume", title_style))
    flowables.append(HRFlowable(width="100%", color=HexColor("#0f3460"),
                                thickness=2, spaceAfter=12))

    # Parse text into sections
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            flowables.append(Spacer(1, 6))
            continue

        # Section headers (lines of = signs or ALL CAPS short lines)
        if stripped.startswith("=" * 10):
            continue
        if stripped.isupper() and len(stripped) < 60:
            flowables.append(Spacer(1, 8))
            flowables.append(HRFlowable(width="100%", color=HexColor("#e0e0e0"),
                                        thickness=0.5, spaceAfter=4))
            flowables.append(Paragraph(stripped, heading_style))
            continue

        # Bullet points
        if stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ").strip()
            flowables.append(Paragraph(f"&bull;  {bullet_text}", body_style))
            continue

        # Numbered items
        if stripped[0].isdigit() and "." in stripped[:4]:
            flowables.append(Paragraph(stripped, body_style))
            continue

        # Normal text
        flowables.append(Paragraph(stripped, body_style))

    doc.build(flowables)
    logger.info("PDF generated: %s", output_path)
    return output_path


def generate_docx(text: str, filename: str = "improved_resume.docx") -> Path:
    """
    Create a professionally-formatted DOCX from plain text.

    Args:
        text:     The improved resume text.
        filename: Output filename.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title = doc.add_heading("Improved Resume", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
        run.font.size = Pt(20)

    doc.add_paragraph("")  # spacer

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Section divider lines
        if stripped.startswith("=" * 10):
            continue

        # Section headers
        if stripped.isupper() and len(stripped) < 60:
            heading = doc.add_heading(stripped.title(), level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(13)
            continue

        # Bullet points
        if stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ").strip()
            para = doc.add_paragraph(bullet_text, style="List Bullet")
            continue

        # Numbered items
        if stripped[0].isdigit() and "." in stripped[:4]:
            doc.add_paragraph(stripped, style="List Number")
            continue

        # Normal paragraph
        doc.add_paragraph(stripped)

    output_path = TEMP_DIR / filename
    doc.save(str(output_path))
    logger.info("DOCX generated: %s", output_path)
    return output_path
